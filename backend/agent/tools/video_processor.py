# KONSPECTO/backend/agent/tools/video_processor.py

import logging
import os
import re
import tempfile
import uuid

from abc import ABC, abstractmethod
from io import BytesIO
from urllib.error import HTTPError

import cv2
import numpy as np

from docx import Document
from docx.shared import Inches
from PIL import Image
from pytubefix import YouTube
from pytubefix.cli import on_progress
from pytubefix.exceptions import RegexMatchError  # Добавляем импорт
from skimage.metrics import structural_similarity as ssim

from app.exceptions import InvalidYouTubeURLException, VideoProcessingError
from app.services.redis_service import RedisService

logger = logging.getLogger("agent.tools.video_processor")


class ImageDifferenceChecker(ABC):
    """
    Абстрактный класс для проверки различий между двумя изображениями.
    """

    @abstractmethod
    def are_images_different(self, img_path1: str, img_path2: str) -> bool:
        """
        Определяет, отличаются ли два изображения.

        :param img_path1: Путь к первому изображению.
        :param img_path2: Путь ко второму изображению.
        :return: True, если изображения отличаются, иначе False.
        """
        pass


class SSIMImageDifferenceChecker(ImageDifferenceChecker):
    """
    Класс для проверки различий между изображениями с использованием SSIM.
    """

    def __init__(self, threshold: float = 0.98):
        self.threshold = threshold

    def are_images_different(self, img_path1: str, img_path2: str) -> bool:
        """
        Сравнивает два изображения с использованием SSIM и определяет,
        отличаются ли они более чем на заданный порог.

        :param img_path1: Путь к первому изображению.
        :param img_path2: Путь ко второму изображению.
        :return: True, если изображения отличаются, иначе False.
        """
        try:
            img1 = Image.open(img_path1).convert("L")  # Преобразование в градации серого
            img2 = Image.open(img_path2).convert("L")

            # Приведение изображений к одному размеру
            if img1.size != img2.size:
                img2 = img2.resize(img1.size)

            arr1 = np.array(img1)
            arr2 = np.array(img2)

            # Вычисление SSIM
            similarity = ssim(arr1, arr2)
            logger.debug(
                f"SSIM similarity between '{img_path1}' and '{img_path2}': {similarity}"
            )

            return similarity < self.threshold

        except Exception as e:
            logger.exception(
                f"Ошибка при сравнении изображений с использованием SSIM: {e}"
            )
            # В случае ошибки считаем изображения различными
            return True


def sanitize_filename(filename: str) -> str:
    """
    Очищает имя файла, заменяя пробелы на подчеркивания и удаляя нежелательные символы.

    :param filename: Исходное имя файла.
    :return: Очищенное имя файла.
    """
    # Заменяем пробелы и другие нежелательные символы на подчеркивания
    filename = re.sub(r"\s+", "_", filename)
    # Удаляем любые символы, кроме букв, цифр, подчеркиваний, дефисов и точек
    filename = re.sub(r"[^\w\-_\.]", "", filename)
    return filename


class VideoToDocxConverter:
    """
    Класс, инкапсулирующий процесс конвертации видео YouTube в DOCX документ.
    """

    def __init__(
        self,
        youtube_url: str,
        redis_service: RedisService,
        difference_checker: ImageDifferenceChecker = None,
        expire_seconds: int = 86400,  # По умолчанию документ истекает через 1 день
    ):
        self.youtube_url = youtube_url
        self.redis_service = redis_service
        self.difference_checker = difference_checker or SSIMImageDifferenceChecker()
        self.expire_seconds = expire_seconds
        self.temp_dir = None
        self.video_path = None
        self.extracted_images = []
        self.unique_key = None
        self.video_title = None

    async def process(self) -> str:
        """
        Выполняет полный процесс конвертации видео и сохранения DOCX документа в Redis.

        :return: Уникальный ключ для доступа к DOCX файлу в Redis.
        """
        try:
            logger.info(f"Начало обработки видео по ссылке: {self.youtube_url}")
            await self.download_video()
            self.extract_images()
            docx_bytes = self.create_docx()
            await self.save_to_redis(docx_bytes)
            return self.unique_key
        except InvalidYouTubeURLException as e:
            logger.error(f"Invalid input: {e.detail}")
            raise e
        except Exception as e:
            logger.exception("Video processing operation failed.")
            raise VideoProcessingError()
        finally:
            self.cleanup()

    async def download_video(self):
        """
        Загрузка видео с YouTube во временную директорию.
        """
        self.temp_dir = tempfile.mkdtemp()
        logger.debug(f"Создана временная директория: {self.temp_dir}")

        try:
            # Загрузка видео с YouTube
            yt = YouTube(self.youtube_url, on_progress_callback=on_progress)
            self.video_title = yt.title
            stream = yt.streams.get_highest_resolution()
            if not stream:
                logger.error("Не удалось найти подходящий поток для загрузки.")
                raise VideoProcessingError(
                    "Не удалось найти подходящий поток для загрузки."
                )

            self.video_path = os.path.join(self.temp_dir, "video.mp4")
            logger.info(f"Загрузка видео: {self.video_title}")
            stream.download(output_path=self.temp_dir, filename="video.mp4")
            logger.info(f"Видео загружено по пути: {self.video_path}")

        except RegexMatchError:
            logger.error(f"Неверный URL YouTube: {self.youtube_url}")
            raise InvalidYouTubeURLException()
        except HTTPError as http_err:
            logger.error(f"HTTP ошибка при загрузке видео: {http_err}")
            raise VideoProcessingError(
                "Доступ к видео запрещен (HTTP 403). Проверьте ссылку или ограничения видео."
            )
        except Exception as e:
            logger.exception(f"Не удалось загрузить видео: {e}")
            raise VideoProcessingError("Не удалось загрузить видео.")

    def extract_images(self):
        """
        Извлечение изображений из видео каждые 5 секунд, удаление схожих изображений.
        """
        cap = cv2.VideoCapture(self.video_path)
        fps = cap.get(cv2.CAP_PROP_FPS)
        if fps == 0:
            logger.error("Не удалось определить FPS видео.")
            raise VideoProcessingError("Не удалось определить FPS видео.")

        frame_interval = int(fps * 5)  # Интервал кадров для каждых 5 секунд
        logger.debug(f"FPS видео: {fps}, интервал кадров: {frame_interval}")

        frame_count = 0
        last_image_path = None

        while True:
            ret, frame = cap.read()
            if not ret:
                break

            if frame_count % frame_interval == 0:
                # Конвертация кадра из BGR (OpenCV) в RGB (Pillow)
                frame_rgb = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
                img = Image.fromarray(frame_rgb)

                img_path = os.path.join(self.temp_dir, f"frame_{frame_count}.png")
                img.save(img_path)
                logger.info(f"Извлечено изображение: {img_path}")

                # Проверка на схожесть с последним сохраненным изображением
                if last_image_path:
                    if self.difference_checker.are_images_different(
                        last_image_path, img_path
                    ):
                        self.extracted_images.append(img_path)
                        last_image_path = img_path
                        logger.debug(
                            f"Изображение {img_path} отличается от предыдущего. Сохранено."
                        )
                    else:
                        logger.debug(
                            f"Изображение {img_path} схоже с предыдущим. Пропущено."
                        )
                        os.remove(img_path)  # Удаляем схожее изображение
                else:
                    self.extracted_images.append(img_path)
                    last_image_path = img_path
                    logger.debug(f"Первое изображение {img_path} сохранено.")

            frame_count += 1

        cap.release()
        logger.info(f"Всего извлечено изображений: {len(self.extracted_images)}")

        if not self.extracted_images:
            logger.warning("Не было извлечено ни одного изображения.")
            raise VideoProcessingError("Не было извлечено ни одного изображения.")

    def create_docx(self) -> bytes:
        """
        Создание DOCX документа с извлеченными изображениями.

        :return: Байтовое представление DOCX документа.
        """
        doc = Document()
        doc.add_heading(self.video_title, 0)

        for img_path in self.extracted_images:
            doc.add_picture(img_path, width=Inches(6))

        # Сохранение DOCX документа в BytesIO
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_bytes = doc_io.getvalue()
        logger.info("DOCX документ создан успешно.")
        return doc_bytes

    async def save_to_redis(self, doc_bytes: bytes):
        """
        Сохранение DOCX документа в Redis с уникальным ключом.

        :param doc_bytes: Байтовое представление DOCX документа.
        """
        self.unique_key = f"docx:{uuid.uuid4()}"

        success = await self.redis_service.set_file(
            self.unique_key, doc_bytes, expire=self.expire_seconds
        )
        if not success:
            logger.error("Не удалось сохранить DOCX файл в Redis.")
            raise VideoProcessingError("Не удалось сохранить документ в хранилище.")

        logger.info(f"DOCX документ сохранен в Redis с ключом: {self.unique_key}")

    def cleanup(self):
        """
        Очистка временных файлов и директорий.
        """
        if self.temp_dir and os.path.exists(self.temp_dir):
            for root, dirs, files in os.walk(self.temp_dir, topdown=False):
                for name in files:
                    os.remove(os.path.join(root, name))
                for name in dirs:
                    os.rmdir(os.path.join(root, name))
            os.rmdir(self.temp_dir)
            logger.debug(f"Удалена временная директория: {self.temp_dir}")


async def youtube_to_docx(
    youtube_url: str,
    redis_service: RedisService,
    difference_checker: ImageDifferenceChecker = None,
    expire_seconds: int = 86400,  # Документ истекает через 1 день по умолчанию
) -> str:
    """
    Обертка для конвертации YouTube видео в DOCX документ и сохранения его в Redis.

    :param youtube_url: Ссылка на YouTube видео.
    :param redis_service: Экземпляр RedisService для взаимодействия с Redis.
    :param difference_checker: Объект для определения различий между изображениями.
    :param expire_seconds: Время в секундах, после которого документ истечет в Redis.
    :return: Уникальный ключ для доступа к DOCX файлу в Redis.
    """
    converter = VideoToDocxConverter(
        youtube_url=youtube_url,
        redis_service=redis_service,
        difference_checker=difference_checker,
        expire_seconds=expire_seconds,
    )
    return await converter.process()
